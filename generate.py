#!/usr/bin/env python3
"""
generate.py
===========
Reads structured content from Word, Excel, PNG figures, and BibTeX,
then generates a publication-ready PDF (MDPI-style) and a formatted
Word document.

Project layout expected
-----------------------
    content/
        text.docx           <- manuscript text (see README for conventions)
        tables.xlsx         <- tables (one sheet per table)
        references.bib.txt  <- BibTeX references

    figures/
        *.png / *.jpg / *.pdf   <- figure files

    template/
        mdpi_style.tex      <- LaTeX preamble (auto-copied if missing)

    output/                 <- generated files land here

Usage
-----
    python generate.py [--pdf] [--word] [--clean]

    --pdf    (default) compile main.tex to PDF
    --word   generate output.docx
    --clean  remove LaTeX auxiliary files after compilation
    --all    generate both PDF and Word
"""

# -- Standard library ----------------------------------------------------------
import argparse
import re
import shutil
import subprocess
import sys
from pathlib import Path

# -- Third-party ---------------------------------------------------------------
try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from openpyxl import load_workbook
except ImportError:
    sys.exit(
        "Missing dependencies. Run:  pip install python-docx openpyxl\n"
        "Or:  pip install -r requirements.txt"
    )

# ------------------------------------------------------------------------------
# CONFIGURATION
# ------------------------------------------------------------------------------

ROOT        = Path(__file__).parent
CONTENT_DIR = ROOT / "content"
FIGURES_DIR = ROOT / "figures"
TEMPLATE    = ROOT / "template" / "mdpi_style.tex"
OUTPUT_DIR  = ROOT / "output"
TEX_OUT     = ROOT / "main.tex"
BIB_OUT     = ROOT / "references.bib"

WORD_FILE   = CONTENT_DIR / "text.docx"
EXCEL_FILE  = CONTENT_DIR / "tables.xlsx"
BIB_FILE    = CONTENT_DIR / "references.bib.txt"

# Row colour map  (Excel column "COLOR" -> LaTeX colour)
COLOR_MAP = {
    "green":  "mcolor!12",
    "blue":   "mdpiblue!12",
    "gray":   "rowgray",
    "yellow": "yellow!20",
    "orange": "orange!15",
    "red":    "red!12",
    "purple": "violet!12",
}

# ------------------------------------------------------------------------------
# HELPERS
# ------------------------------------------------------------------------------

def ensure_output():
    OUTPUT_DIR.mkdir(exist_ok=True)


def esc(text: str) -> str:
    """Escape LaTeX special characters in plain text."""
    if not text:
        return ""
    text = str(text)
    replacements = [
        ("\\", r"\textbackslash{}"),
        ("&",  r"\&"),
        ("%",  r"\%"),
        ("$",  r"\$"),
        ("#",  r"\#"),
        ("_",  r"\_"),
        ("{",  r"\{"),
        ("}",  r"\}"),
        ("~",  r"\textasciitilde{}"),
        ("^",  r"\textasciicircum{}"),
    ]
    for old, new in replacements:
        if old == "\\":
            text = text.replace(old, new)
        else:
            text = re.sub(r"(?<!\\)" + re.escape(old), new, text)
    # Typography
    text = text.replace("\n", r" \\ ")
    text = text.replace("->", r"$\to$")
    text = text.replace("≥", r"$\geq$")
    text = text.replace("≤", r"$\leq$")
    text = text.replace("–", "--")
    text = text.replace("—", "---")
    text = text.replace("≈", r"$\approx$")
    # Greek letters and math symbols
    greek = [("λ","\\lambda"),("β","\\beta"),("α","\\alpha"),
             ("γ","\\gamma"),("δ","\\delta"),("σ","\\sigma"),
             ("μ","\\mu"),("π","\\pi"),("Δ","\\Delta"),
             ("±","\\pm"),("×","\\times"),("÷","\\div")]
    for ch, cmd in greek:
        text = text.replace(ch, f"${cmd}$")
    return text


def run_to_latex(run) -> str:
    """Convert a python-docx Run to LaTeX, preserving bold/italic."""
    txt = esc(run.text)
    if not txt.strip():
        return txt
    if run.bold and run.italic:
        return rf"\textbf{{\textit{{{txt}}}}}"
    if run.bold:
        return rf"\textbf{{{txt}}}"
    if run.italic:
        return rf"\textit{{{txt}}}"
    return txt


# ------------------------------------------------------------------------------
# BIBTEX PARSER
# ------------------------------------------------------------------------------

def process_bibtex(src: Path, dst: Path) -> list[str]:
    """
    Copy the BibTeX file to output path and return list of citation keys.

    The source file uses the .bib.txt extension so it opens in any text
    editor without BibTeX associations. We write it as a plain .bib for
    LaTeX to consume.
    """
    if not src.exists():
        print(f"  [warn] BibTeX file not found: {src}")
        dst.write_text("% No references provided\n", encoding="utf-8")
        return []

    raw = src.read_text(encoding="utf-8")
    dst.write_text(raw, encoding="utf-8")

    keys = re.findall(r"@\w+\s*\{\s*([^,\s]+)", raw)
    print(f"  -> BibTeX: {len(keys)} entries -> {dst.name}")
    return keys


# ------------------------------------------------------------------------------
# WORD READER
# ------------------------------------------------------------------------------

MARKER_TABLE  = re.compile(r"^TABLE\s*:\s*(.+)$", re.IGNORECASE)
MARKER_FIGURE = re.compile(r"^FIGURE\s*:\s*(.+)$", re.IGNORECASE)


def para_to_latex(para) -> str:
    """Convert a single Word paragraph to a LaTeX string."""
    style = para.style.name
    raw   = para.text.strip()

    if not raw:
        return ""

    # -- Markers --------------------------------------------------------------
    m = MARKER_TABLE.match(raw)
    if m:
        return f"%%TABLE:{m.group(1).strip()}%%"

    m = MARKER_FIGURE.match(raw)
    if m:
        return f"%%FIGURE:{m.group(1).strip()}%%"

    # -- Run-level formatting --------------------------------------------------
    runs_tex = "".join(run_to_latex(r) for r in para.runs)

    # -- Headings --------------------------------------------------------------
    if "Heading 1" in style or style == "Title":
        runs_tex = re.sub(r"^\d+[\.\)]\s*", "", runs_tex)
        return f"\n\\section{{{runs_tex}}}\n"
    if "Heading 2" in style:
        return f"\n\\subsection{{{runs_tex}}}\n"
    if "Heading 3" in style:
        return f"\n\\subsubsection{{{runs_tex}}}\n"

    # -- Lists -----------------------------------------------------------------
    if style.startswith("List Number"):
        return f"  \\item {runs_tex}"
    if style.startswith("List Bullet"):
        return f"  \\item {runs_tex}"

    # -- Quote -----------------------------------------------------------------
    if style == "Quote" or style == "Intense Quote":
        return f"\\begin{{quote}}\n{runs_tex}\n\\end{{quote}}"

    return runs_tex


def read_word(path: Path) -> dict:
    """
    Parse the Word document and return a structured dict:
      {
        "title": str,
        "short_title": str,
        "authors": str,
        "short_authors": str,
        "affiliations": str,
        "journal": str,
        "date": str,
        "abstract": str,
        "keywords": str,
        "blocks": list[str],
        "metadata": dict
      }

    The first few paragraphs with style 'Title', 'Subtitle', or matching
    meta-keys are parsed as metadata. Everything else becomes blocks.
    """
    doc = Document(path)

    meta = {
        "title":        "Document Title",
        "short_title":  "",
        "authors":      "Author Name",
        "short_authors":"Author et al.",
        "affiliations": "",
        "journal":      "",
        "date":         "",
        "abstract":     "",
        "keywords":     "",
        "doi":          "",
    }

    blocks  = []
    in_list = False
    list_env = None

    for para in doc.paragraphs:
        style = para.style.name
        raw   = para.text.strip()

        # -- Metadata lines (KEY: value) ---------------------------------------
        meta_match = re.match(r"^(TITLE|SHORT_TITLE|AUTHORS|SHORT_AUTHORS|"
                              r"AFFILIATIONS|JOURNAL|DATE|ABSTRACT|KEYWORDS|"
                              r"DOI)\s*:\s*(.+)$", raw, re.IGNORECASE)
        if meta_match:
            key = meta_match.group(1).lower()
            val = meta_match.group(2).strip()
            meta[key] = val
            continue

        # -- Handle list open/close --------------------------------------------
        is_num    = style.startswith("List Number")
        is_bullet = style.startswith("List Bullet")
        is_list   = is_num or is_bullet

        if is_list and not in_list:
            in_list  = True
            list_env = "enumerate" if is_num else "itemize"
            blocks.append(rf"\begin{{{list_env}}}")

        if not is_list and in_list:
            blocks.append(rf"\end{{{list_env}}}")
            in_list  = False
            list_env = None

        line = para_to_latex(para)
        if line:
            blocks.append(line)

    if in_list:
        blocks.append(rf"\end{{{list_env}}}")

    # Fall back: first Heading 1 or Title paragraph
    for para in doc.paragraphs:
        if para.style.name in ("Title", "Heading 1") and meta["title"] == "Document Title":
            meta["title"] = para.text.strip()
            break

    if not meta["short_title"]:
        meta["short_title"] = meta["title"][:60]

    return {"meta": meta, "blocks": blocks}


# ------------------------------------------------------------------------------
# EXCEL READER
# ------------------------------------------------------------------------------

def read_excel(path: Path) -> dict[str, str]:
    """
    Return {sheet_name: latex_table_string}.

    Sheet layout:
      Row 1  - table caption (first non-empty cell)
      Row 2  - column headers
      Row 3+ - data rows
      Last column named "COLOR" (case-insensitive) -> row shading
    """
    wb     = load_workbook(path)
    tables = {}

    for ws in wb.worksheets:
        if ws.title.upper() in ("README", "INSTRUCTIONS", "INSTRUCOES"):
            continue

        rows = list(ws.iter_rows(values_only=True))
        if len(rows) < 3:
            continue

        caption = esc(str(rows[0][0])) if rows[0][0] else ws.title
        headers = [str(c).strip() if c else "" for c in rows[1]]

        # Detect optional COLOR column
        has_color = headers[-1].upper() == "COLOR"
        data_hdrs = headers[:-1] if has_color else headers
        n = len(data_hdrs)

        # Auto column spec
        col_spec = _auto_colspec(n)

        hdr_row = " & ".join(rf"\textbf{{{esc(h)}}}" for h in data_hdrs)

        data_rows_tex = []
        for row in rows[2:]:
            if all(c is None for c in row):
                continue
            cells = list(row)
            color_tex = ""
            if has_color and cells:
                color_raw = str(cells[-1]).strip().lower() if cells[-1] else ""
                color_tex = COLOR_MAP.get(color_raw, "")
                cells     = cells[:-1]
            while len(cells) < n:
                cells.append("")
            cell_tex = " & ".join(
                esc(str(c)) if c is not None else "" for c in cells[:n]
            )
            rc = f"  \\rowcolor{{{color_tex}}}\n" if color_tex else ""
            data_rows_tex.append(f"{rc}  {cell_tex}\\\\")

        latex = (
            f"\\begin{{table}}[H]\n"
            f"\\caption{{{caption}\\label{{tab:{ws.title}}}}}\n"
            f"\\begin{{tabularx}}{{\\textwidth}}{{{col_spec}}}\n"
            f"\\toprule\n"
            f"\\rowcolor{{rowgray}}\n"
            f"  {hdr_row}\\\\\n"
            f"\\midrule\n"
            + "\n".join(data_rows_tex) + "\n"
            f"\\bottomrule\n"
            f"\\end{{tabularx}}\n"
            f"\\end{{table}}"
        )
        tables[ws.title] = latex

    return tables


def _auto_colspec(n: int) -> str:
    """Generate a reasonable tabularX column specification."""
    specs = {
        1: "X",
        2: "L{4cm} X",
        3: "L{3cm} L{4cm} X",
        4: "L{2.5cm} L{3cm} L{3.5cm} X",
        5: "L{1.8cm} L{2.5cm} L{2.5cm} L{3cm} X",
        6: "L{1.5cm} L{2cm} L{2cm} L{2cm} L{2.5cm} X",
        7: "L{1.3cm} L{1.8cm} L{1.8cm} C{1.2cm} C{1.2cm} C{1.2cm} X",
    }
    return specs.get(n, " ".join(["L{2cm}"] * (n - 1)) + " X")


# ------------------------------------------------------------------------------
# FIGURE HANDLER
# ------------------------------------------------------------------------------

def figure_latex(spec: str, fig_dir: Path) -> str:
    """
    Parse a figure spec string and return a LaTeX figure environment.

    Spec format:  filename.png | Caption text | width_fraction
    Examples:
        map.png | Study area map | 0.80
        chart.png | Results by category
        photo.png
    """
    parts    = [p.strip() for p in spec.split("|")]
    filename = parts[0]
    caption  = esc(parts[1]) if len(parts) > 1 else esc(filename)
    try:
        width = float(parts[2]) if len(parts) > 2 else 0.85
        width = max(0.1, min(1.0, width))
    except ValueError:
        width = 0.85

    fpath = fig_dir / filename
    if not fpath.exists():
        return (
            f"\n% [FIGURE NOT FOUND: {filename}  -  "
            f"place file in figures/ directory]\n"
        )

    label = re.sub(r"[^a-zA-Z0-9]", "_", filename.rsplit(".", 1)[0])
    return (
        f"\n\\begin{{figure}}[H]\n"
        f"\\centering\n"
        f"\\includegraphics[width={width:.2f}\\linewidth]{{figures/{filename}}}\n"
        f"\\caption{{{caption}\\label{{fig:{label}}}}}\n"
        f"\\end{{figure}}\n"
    )


# ------------------------------------------------------------------------------
# LATEX BUILDER
# ------------------------------------------------------------------------------

def build_latex(parsed: dict, tables: dict, fig_dir: Path,
                bib_keys: list[str]) -> str:
    """Assemble the complete LaTeX document string."""

    meta   = parsed["meta"]
    blocks = parsed["blocks"]

    # -- Preamble from template file -------------------------------------------
    preamble = TEMPLATE.read_text(encoding="utf-8")
    # Strip the \documentclass declaration; we rewrite it below.
    preamble = re.sub(r"\\documentclass.*?\n", "", preamble, count=1)

    lines = [
        r"\documentclass[11pt, a4paper]{article}",
        preamble,
        "",
        rf"\shorttitle{{{esc(meta['short_title'])}}}",
        rf"\shortauthor{{{esc(meta['short_authors'])}}}",
        rf"\journalname{{{esc(meta['journal'])}}}",
        r"\graphicspath{{figures/}}",
        "",
        r"\begin{document}",
        "",
    ]

    # -- Title block -----------------------------------------------------------
    lines += [
        rf"\title{{{esc(meta['title'])}}}",
        rf"\author{{{esc(meta['authors'])}}}",
        rf"\date{{{esc(meta['date'])}}}",
        r"\maketitle",
        r"\thispagestyle{fancy}",
        "",
    ]

    # -- Affiliations ----------------------------------------------------------
    if meta["affiliations"]:
        lines += [
            rf"{{\small\itshape {esc(meta['affiliations'])}}}\\[4pt]",
            "",
        ]

    # -- Abstract --------------------------------------------------------------
    if meta["abstract"]:
        lines += [
            r"\begin{abstractbox}",
            esc(meta["abstract"]),
            r"\end{abstractbox}",
            "",
        ]

    # -- Keywords --------------------------------------------------------------
    if meta["keywords"]:
        lines += [
            rf"\noindent\textbf{{Keywords:}} {esc(meta['keywords'])}",
            "",
            r"\noindent\textcolor{mcolor}{\rule{\linewidth}{0.4pt}}",
            "",
        ]

    # -- Body ------------------------------------------------------------------
    i = 0
    while i < len(blocks):
        blk = blocks[i]

        # Table marker
        if blk.startswith("%%TABLE:") and blk.endswith("%%"):
            name = blk[8:-2].strip()
            if name in tables:
                lines.append("\n" + tables[name] + "\n")
            else:
                avail = ", ".join(tables) or "none"
                lines.append(
                    f"\n% [TABLE NOT FOUND: '{name}'  -  "
                    f"available sheets: {avail}]\n"
                )
            i += 1
            continue

        # Figure marker
        if blk.startswith("%%FIGURE:") and blk.endswith("%%"):
            spec = blk[9:-2].strip()
            lines.append(figure_latex(spec, fig_dir))
            i += 1
            continue

        lines.append(blk)
        i += 1

    # -- Bibliography ----------------------------------------------------------
    lines += [
        "",
        r"\bibliographystyle{unsrtnat}",
        r"\bibliography{references}",
        "",
        r"\end{document}",
        "",
    ]

    return "\n".join(lines)


# ------------------------------------------------------------------------------
# WORD BUILDER
# ------------------------------------------------------------------------------

def build_word(parsed: dict, tables_raw: dict,
               fig_dir: Path) -> Document:
    """Generate a formatted Word document from the parsed content."""

    doc  = Document()
    meta = parsed["meta"]

    # -- Page setup ------------------------------------------------------------
    section = doc.sections[0]
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin   = Inches(1.2)
    section.right_margin  = Inches(1.0)

    # -- Base font -------------------------------------------------------------
    style = doc.styles["Normal"]
    style.font.name = "Palatino Linotype"
    style.font.size = Pt(11)

    GREEN = RGBColor(26, 107, 58)
    BLUE  = RGBColor(0, 103, 172)

    def add_heading(text, level=1, color=GREEN):
        h = doc.add_heading(text, level=level)
        for run in h.runs:
            run.font.color.rgb = color
            run.font.name = "Arial"
        return h

    def add_para(text, bold=False, italic=False, size=11, align=None):
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.bold   = bold
        r.italic = italic
        r.font.size = Pt(size)
        r.font.name = "Palatino Linotype"
        if align:
            p.alignment = align
        return p

    # -- Title -----------------------------------------------------------------
    tp = doc.add_paragraph()
    tr = tp.add_run(meta["title"])
    tr.bold      = True
    tr.font.size = Pt(16)
    tr.font.color.rgb = GREEN
    tr.font.name = "Arial"

    if meta["authors"]:
        ap = doc.add_paragraph()
        ar = ap.add_run(meta["authors"])
        ar.font.size = Pt(11)
        ar.font.name = "Palatino Linotype"

    if meta["affiliations"]:
        af = doc.add_paragraph()
        afr = af.add_run(meta["affiliations"])
        afr.font.size   = Pt(9)
        afr.font.italic = True
        afr.font.name   = "Palatino Linotype"

    if meta["date"]:
        add_para(meta["date"], italic=True, size=10)

    doc.add_paragraph()  # spacer

    # -- Abstract -------------------------------------------------------------
    if meta["abstract"]:
        add_heading("Abstract", level=1)
        add_para(meta["abstract"])

    if meta["keywords"]:
        kp = doc.add_paragraph()
        kr = kp.add_run("Keywords: ")
        kr.bold = True
        kr.font.name = "Palatino Linotype"
        kr2 = kp.add_run(meta["keywords"])
        kr2.font.name = "Palatino Linotype"

    doc.add_paragraph()

    # -- Body blocks -----------------------------------------------------------
    in_list = False
    list_type = None

    for blk in parsed["blocks"]:

        # Table marker; simplified plain table.
        if blk.startswith("%%TABLE:") and blk.endswith("%%"):
            _add_word_table(doc, blk[8:-2].strip(), tables_raw, GREEN)
            continue

        # Figure marker
        if blk.startswith("%%FIGURE:") and blk.endswith("%%"):
            _add_word_figure(doc, blk[9:-2].strip(), fig_dir)
            continue

        # Section headings
        if blk.startswith("\n\\section{"):
            title = re.search(r"\\section\{(.+)\}", blk)
            if title:
                add_heading(title.group(1), level=1)
            continue
        if blk.startswith("\n\\subsection{"):
            title = re.search(r"\\subsection\{(.+)\}", blk)
            if title:
                add_heading(title.group(1), level=2, color=BLUE)
            continue
        if blk.startswith("\n\\subsubsection{"):
            title = re.search(r"\\subsubsection\{(.+)\}", blk)
            if title:
                add_heading(title.group(1), level=3, color=BLUE)
            continue

        # Lists
        if r"\begin{enumerate}" in blk or r"\begin{itemize}" in blk:
            in_list  = True
            list_type = "num" if "enumerate" in blk else "bullet"
            continue
        if r"\end{enumerate}" in blk or r"\end{itemize}" in blk:
            in_list  = False
            list_type = None
            continue
        if blk.startswith("  \\item "):
            text = blk[8:]
            text = _strip_latex_cmds(text)
            style_name = "List Number" if list_type == "num" else "List Bullet"
            p = doc.add_paragraph(text, style=style_name)
            p.style.font.name = "Palatino Linotype"
            continue

        # Normal paragraph
        if blk.strip():
            text = _strip_latex_cmds(blk)
            add_para(text)

    return doc


def _strip_latex_cmds(text: str) -> str:
    """Remove common LaTeX commands for plain Word output."""
    text = re.sub(r"\\textbf\{([^}]+)\}", r"\1", text)
    text = re.sub(r"\\textit\{([^}]+)\}", r"\1", text)
    text = re.sub(r"\\emph\{([^}]+)\}", r"\1", text)
    text = re.sub(r"\\texttt\{([^}]+)\}", r"\1", text)
    text = re.sub(r"\$\\to\$", "->", text)
    text = re.sub(r"\$\\geq\$", "=", text)
    text = re.sub(r"\$\\leq\$", "=", text)
    text = re.sub(r"\$\\approx\$", "≈", text)
    text = re.sub(r"\\&", "&", text)
    text = re.sub(r"\\%", "%", text)
    text = re.sub(r"\\_", "_", text)
    text = re.sub(r"\\\$", "$", text)
    text = re.sub(r"---", "—", text)
    text = re.sub(r"--", "–", text)
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def _add_word_table(doc: Document, name: str,
                    tables_raw: dict, header_color):
    """Add a simplified table to the Word document."""
    from openpyxl import load_workbook
    if not EXCEL_FILE.exists():
        return
    wb = load_workbook(EXCEL_FILE)
    if name not in wb.sheetnames:
        doc.add_paragraph(f"[Table not found: {name}]")
        return
    ws   = wb[name]
    rows = list(ws.iter_rows(values_only=True))
    if len(rows) < 2:
        return

    caption = str(rows[0][0]) if rows[0][0] else name
    p = doc.add_paragraph()
    r = p.add_run(f"Table: {caption}")
    r.bold = True; r.font.name = "Arial"; r.font.size = Pt(10)

    headers = [str(c) if c else "" for c in rows[1]]
    has_col = headers[-1].upper() == "COLOR"
    if has_col:
        headers = headers[:-1]

    n = len(headers)
    tbl = doc.add_table(rows=1, cols=n)
    tbl.style = "Table Grid"

    # Header row
    hdr_cells = tbl.rows[0].cells
    for j, h in enumerate(headers):
        hdr_cells[j].text = h
        for run in hdr_cells[j].paragraphs[0].runs:
            run.bold = True
            run.font.name = "Arial"
            run.font.size = Pt(9)

    # Data rows
    for row in rows[2:]:
        if all(c is None for c in row):
            continue
        cells = list(row)
        if has_col:
            cells = cells[:-1]
        while len(cells) < n:
            cells.append("")
        row_cells = tbl.add_row().cells
        for j, c in enumerate(cells[:n]):
            row_cells[j].text = str(c) if c is not None else ""
            for run in row_cells[j].paragraphs[0].runs:
                run.font.name = "Palatino Linotype"
                run.font.size = Pt(9)

    doc.add_paragraph()


def _add_word_figure(doc: Document, spec: str, fig_dir: Path):
    """Add a figure to the Word document."""
    parts    = [p.strip() for p in spec.split("|")]
    filename = parts[0]
    caption  = parts[1] if len(parts) > 1 else filename
    try:
        width = float(parts[2]) if len(parts) > 2 else 0.85
    except ValueError:
        width = 0.85

    fpath = fig_dir / filename
    if fpath.exists() and fpath.suffix.lower() in (".png", ".jpg", ".jpeg"):
        try:
            doc.add_picture(str(fpath), width=Inches(6.0 * width))
            last = doc.paragraphs[-1]
            last.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception:
            pass

    cp = doc.add_paragraph()
    cr = cp.add_run(f"Figure: {caption}")
    cr.italic    = True
    cr.font.size = Pt(9)
    cr.font.name = "Palatino Linotype"
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()


# ------------------------------------------------------------------------------
# COMPILER
# ------------------------------------------------------------------------------

def compile_pdf(clean: bool = False) -> bool:
    """Run pdflatex twice to resolve cross-references."""
    cmd = ["pdflatex", "-interaction=nonstopmode",
           "-output-directory", str(ROOT), str(TEX_OUT)]

    for run_n in [1, 2]:
        print(f"  -> pdflatex pass {run_n}/2 ...", end=" ", flush=True)
        result = subprocess.run(cmd, capture_output=True, text=True, cwd=ROOT)
        if result.returncode != 0:
            # Find first error line
            for line in result.stdout.splitlines():
                if line.startswith("!"):
                    print(f"\n  [error] {line}")
                    break
            print("FAILED")
            print("  Run 'pdflatex main.tex' manually to see full log.")
            return False
        print("OK")

    pdf_src = ROOT / "main.pdf"
    pdf_dst = OUTPUT_DIR / "output.pdf"
    if pdf_src.exists():
        shutil.copy2(pdf_src, pdf_dst)
        print(f"  -> PDF saved: {pdf_dst}")

    if clean:
        _clean_aux()

    return True


def _clean_aux():
    exts = [".aux", ".log", ".out", ".toc", ".fls",
            ".fdb_latexmk", ".synctex.gz", ".bbl", ".blg"]
    for ext in exts:
        f = ROOT / f"main{ext}"
        if f.exists():
            f.unlink()
    print("  -> Auxiliary files removed.")


# ------------------------------------------------------------------------------
# MAIN
# ------------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(
        description="Generate PDF and/or Word from Word + Excel + figures + BibTeX"
    )
    parser.add_argument("--pdf",   action="store_true", help="Generate PDF (default)")
    parser.add_argument("--word",  action="store_true", help="Generate Word document")
    parser.add_argument("--all",   action="store_true", help="Generate both PDF and Word")
    parser.add_argument("--clean", action="store_true", help="Remove LaTeX aux files after build")
    args = parser.parse_args()

    do_pdf  = args.pdf or args.all or not (args.word)
    do_word = args.word or args.all

    ensure_output()

    # -- Check inputs ----------------------------------------------------------
    for f in [WORD_FILE, EXCEL_FILE]:
        if not f.exists():
            sys.exit(f"[error] Required file not found: {f}\n"
                     f"       See README.md for how to set up content/")

    pngs = sorted(FIGURES_DIR.glob("*"))
    pngs = [p for p in pngs if p.suffix.lower() in (".png", ".jpg", ".jpeg", ".pdf")]
    print(f"[info] Figures found: {len(pngs)}")
    for p in pngs:
        print(f"       - {p.name}")

    # -- Parse content ---------------------------------------------------------
    print("\n[step 1/4] Parsing Word document ...")
    parsed = read_word(WORD_FILE)
    print(f"           Title: {parsed['meta']['title'][:60]}")

    print("[step 2/4] Parsing Excel tables ...")
    tables = read_excel(EXCEL_FILE)
    print(f"           Sheets: {', '.join(tables) or 'none'}")

    print("[step 3/4] Processing BibTeX ...")
    bib_keys = process_bibtex(BIB_FILE, BIB_OUT)

    # -- Generate LaTeX --------------------------------------------------------
    print("[step 4/4] Assembling main.tex ...")
    tex = build_latex(parsed, tables, FIGURES_DIR, bib_keys)
    TEX_OUT.write_text(tex, encoding="utf-8")
    print(f"           Written: {TEX_OUT}")

    # -- PDF -------------------------------------------------------------------
    if do_pdf:
        print("\n[PDF] Compiling ...")
        ok = compile_pdf(clean=args.clean)
        if not ok:
            print("[PDF] Build failed. Check main.log for details.")

    # -- Word ------------------------------------------------------------------
    if do_word:
        print("\n[Word] Building output.docx ...")
        word_doc = build_word(parsed, tables, FIGURES_DIR)
        word_dst = OUTPUT_DIR / "output.docx"
        word_doc.save(str(word_dst))
        print(f"[Word] Saved: {word_dst}")

    print("\n[OK]  Done. Check the output/ folder.")


if __name__ == "__main__":
    main()
